from docx import Document
import re,time, os

def extract_and_format_text(docx_path):
    txt_path="output.txt"
    print("Reading file...")    
    doc = Document(docx_path)      
    with open(txt_path, 'w', encoding='utf-8') as txt_file:        
        for para in doc.paragraphs:           
            all_bold = all(run.bold for run in para.runs)         
            formatted_text = ''.join(run.text for run in para.runs)                        
            if all_bold:
                formatted_text = f'###{formatted_text}###'          
            
            txt_file.write(formatted_text + '\n')

def parse_text_file():
    file_path="output.txt"   
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()        
    pattern = re.compile(r'###\s*(.*?)\s*###(.*?)(?=###|$)', re.DOTALL)
    matches = pattern.findall(content)
    results = []
    for title, paragraph in matches:
        results.append("## "+title.strip() +"\n\n"+paragraph.strip()+"\n\n\n")    
    return results

def concatenate_strings(strings, max_length=3000):
    concatenated_strings = []
    current_string = ""
    for string in strings:
        if len(current_string) + len(string) <= max_length:
            current_string += string
        else:
            concatenated_strings.append(current_string)
            current_string = string

    if current_string:
        concatenated_strings.append(current_string)

    return concatenated_strings


def summarize_paragraph(text_to_summarize,model): 
    time.sleep(2)    #OPTIONAL: WAIT BETWEEN REQUESTS IF YOU HAVE LIMITED REQUESTS ON YOUR ACCOUNT
    basic_prompt="\nRiassumi il testo di seguito e non ciò che non è contenuto al suo interno. Fornisci solamente il riassunto senza introduzioni e altri commenti. Dai un titolo al riassunto e mettilo all'inizio della risposta."         
    basic_prompt+="\n"+text_to_summarize
    response = model.generate_content(basic_prompt)
    return response.text

def summarize_all(chunks,model):
    print("File read correctly.\n")    
    total_chunks=len(chunks)
    output_file="output.txt"
    with open(output_file, "w") as file:
        file.write("")
    whole_summary=""
    i=0
    errors=[]
    for chunk in chunks:
        with open(f"output/chunk{i}.txt", "w", encoding='utf-8') as f:
            f.write(chunk)
        i+=1
        percentage=int(i/total_chunks*100)
        print(f"Summarizing in progress... {percentage} %       ",end="\r")
        try:
            paragraph_summary=summarize_paragraph(chunk,model)
            with open(output_file, 'a', encoding='utf-8') as file:
                file.write(paragraph_summary+"\n\n\n")
        except:
            errors.append(i)
    print("Error: could not summarize chunks: "+ str(errors))    
    print("Summarization completed! Conversion in progress, please wait...")


def txt_to_docx(docx_file):
    txt_file="output.txt"
    with open(txt_file, 'r', encoding='utf-8') as f:
        text = f.read()
    document = Document()
    paragraph = document.add_paragraph()    
    for line in text.splitlines():
      if line:  
          paragraph.add_run(line)
      else:
          paragraph = document.add_paragraph()
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("##"):
            paragraph.text=paragraph.text.replace("## ", "")
            for run in paragraph.runs:
                run.bold = True
    document.save(docx_file)
    print("Conversion completed.")
    
    def clean_all():
        os.remove("output.txt")
        
